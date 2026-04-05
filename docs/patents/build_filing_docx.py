#!/usr/bin/env python
"""
build_filing_docx.py
====================

Builds USPTO non-provisional utility patent application DOCX files.

Produces ONE consolidated DOCX per patent family containing all sections
in USPTO-required order:
  1. Title of Invention
  2. Cross-Reference to Related Applications
  3. Field of the Invention
  4. Background
  5. Summary of the Invention
  6. Brief Description of the Drawings
  7. Detailed Description of the Drawings
  8. Detailed Description (full specification)
  9. Claims
  10. Abstract
  11. Drawing Sheets (embedded)

Also produces a separate Drawings-only DOCX (one figure per page).

Usage:
    python build_filing_docx.py --family Family_A_NeuroPRIN
    python build_filing_docx.py --family all

Dependencies:
    pip install python-docx
"""

import argparse
import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


PATENTS_DIR = Path(__file__).parent
FAMILIES = [
    "Family_A_NeuroPRIN",
    "Family_B_ALG",
    "Family_C_NeurOS",
    "Family_D_Bridges",
    "Family_E_Applied_Energetics",
    "Family_F_SkillChain",
    "Family_G_Terra_Unita",
]

TITLES = {
    "Family_A_NeuroPRIN": "Systems and Methods for Runtime Trust-Based Authority Governance in Autonomous Systems with Neural Trust Inference and Tamper-Evident Audit",
    "Family_B_ALG": "Runtime Trust-Based Governance System for Safety-Critical Adaptive Systems",
    "Family_C_NeurOS": "AI-Native Operating System with Embedded Governance, Cognitive Memory, and Self-Recursive Autonomy",
    "Family_D_Bridges": "Systems and Methods for Cross-Subsystem Governance Coordination in Heterogeneous Autonomous Computing Environments",
    "Family_E_Applied_Energetics": "Systems and Methods for Cross-Domain Resonance Analysis and Anti-Fragile Resource Allocation Using Coupled Oscillator Models",
    "Family_F_SkillChain": "Systems and Methods for Trust-Weighted Byzantine Fault Tolerant Consensus, Decentralized Skill Validation, and Governed Skill Interoperability in Multi-Agent Computing Networks",
    "Family_G_Terra_Unita": "Production-Backed Digital Currency with Federated Governance and Trust-Based Monetary Policy",
}


def setup_doc():
    """Create and configure a new document with USPTO formatting."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # Add page numbering in footer
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar1)
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    run._r.append(instrText)
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar2)

    return doc


def add_heading(doc, text, level=1):
    """Add a properly formatted heading."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"


_para_counter = [0]  # mutable counter for paragraph numbering


def reset_para_counter():
    _para_counter[0] = 0


def next_para_number():
    _para_counter[0] += 1
    return f"[{_para_counter[0]:04d}]"


def add_para(doc, text, bold=False, italic=False, align=None, size=12, numbered=False):
    """Add a paragraph with optional formatting and USPTO paragraph numbering."""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if numbered:
        prefix = next_para_number()
        run = p.add_run(f"{prefix}  ")
        run.bold = False
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    return p


def add_formatted_text(paragraph, text):
    """Add text with **bold** and *italic* markdown inline formatting."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            subparts = re.split(r"(\*.*?\*)", part)
            for sp in subparts:
                if sp.startswith("*") and sp.endswith("*") and not sp.startswith("**"):
                    run = paragraph.add_run(sp[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(sp)
        if paragraph.runs:
            paragraph.runs[-1].font.name = "Times New Roman"
            paragraph.runs[-1].font.size = Pt(12)


def md_to_docx(doc, md_text, skip_title=True, skip_sections=None, numbered=False):
    """
    Convert markdown text to docx paragraphs.
    skip_sections: list of section heading patterns to skip
    numbered: if True, add [0001] paragraph numbering (USPTO spec requirement)
    """
    skip_sections = skip_sections or []
    lines = md_text.split("\n")
    i = 0
    skipping = False

    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines
        if not line:
            i += 1
            skipping = False
            continue

        # Skip metadata lines
        if any(line.startswith(prefix) for prefix in [
            "*Classification:", "*Applicant:", "*Prepared:", "*Date:",
            "**DRAFT", "**Status:", "**Classification:", "**Date:"
        ]):
            i += 1
            continue

        # Horizontal rules
        if line.startswith("---"):
            i += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip("*# ")

            # Skip title line if requested
            if skip_title and level == 1 and i < 3:
                i += 1
                continue

            # Check if this section should be skipped
            should_skip = False
            for pattern in skip_sections:
                if pattern.lower() in text.lower():
                    should_skip = True
                    break

            if should_skip:
                # Skip until next same-level or higher heading
                i += 1
                while i < len(lines):
                    next_match = re.match(r"^(#{1," + str(level) + r"})\s+", lines[i])
                    if next_match:
                        break
                    i += 1
                continue

            add_heading(doc, text, level=min(level, 4))
            i += 1
            continue

        # Bold-only lines
        bold_match = re.match(r"^\*\*(.+)\*\*\s*$", line)
        if bold_match and ":" not in bold_match.group(1):
            add_para(doc, bold_match.group(1), bold=True)
            i += 1
            continue

        # Key-value bold lines like **Title:** Value
        kv_match = re.match(r"^\*\*(.+?):\*\*\s*(.*)$", line)
        if kv_match:
            p = doc.add_paragraph()
            run = p.add_run(f"{kv_match.group(1)}: ")
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run = p.add_run(kv_match.group(2))
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            i += 1
            continue

        # Bullet lists
        list_match = re.match(r"^(\s*)[-*]\s+(.+)$", line)
        if list_match:
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_text(p, list_match.group(2))
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r"^(\s*)\d+[.)]\s+(.+)$", line)
        if num_match:
            p = doc.add_paragraph(style="List Number")
            add_formatted_text(p, num_match.group(2))
            i += 1
            continue

        # Tables - convert to simple text
        if line.startswith("|"):
            # Skip separator rows
            if re.match(r"^\|[\s\-|]+\|$", line):
                i += 1
                continue
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if cells:
                p = doc.add_paragraph()
                add_formatted_text(p, "  |  ".join(cells))
            i += 1
            continue

        # Code blocks
        if line.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(10)
            i += 1
            continue

        # Regular paragraph
        para_lines = [line]
        i += 1
        while i < len(lines):
            nl = lines[i].rstrip()
            if (not nl or nl.startswith("#") or nl.startswith("---") or
                nl.startswith("- ") or nl.startswith("* ") or nl.startswith("|") or
                nl.startswith("```") or re.match(r"^\*\*.*\*\*\s*$", nl) or
                re.match(r"^\d+[.)]\s+", nl)):
                break
            para_lines.append(nl)
            i += 1

        text = " ".join(para_lines)
        p = doc.add_paragraph()
        if numbered:
            prefix_run = p.add_run(f"{next_para_number()}  ")
            prefix_run.font.name = "Times New Roman"
            prefix_run.font.size = Pt(12)
        add_formatted_text(p, text)


def extract_abstract(md_text):
    """Extract abstract section from markdown."""
    match = re.search(r"##\s*ABSTRACT\s*\n(.*?)(?=\n##|\Z)", md_text, re.DOTALL | re.IGNORECASE)
    if match:
        abstract = match.group(1).strip()
        # Remove markdown formatting
        abstract = re.sub(r"\*\*(.+?)\*\*", r"\1", abstract)
        abstract = re.sub(r"\*(.+?)\*", r"\1", abstract)
        return abstract
    return None


def build_nonprovisional_docx(family_name, output_dir):
    """Build a complete non-provisional utility patent application DOCX."""
    family_dir = PATENTS_DIR / family_name
    title = TITLES.get(family_name, "TITLE NOT SET")

    # Read all source documents
    spec_text = (family_dir / "Specification.md").read_text(encoding="utf-8") if (family_dir / "Specification.md").exists() else ""
    prov_text = (family_dir / "Provisional.md").read_text(encoding="utf-8") if (family_dir / "Provisional.md").exists() else ""
    claims_text = (family_dir / "Claims.md").read_text(encoding="utf-8") if (family_dir / "Claims.md").exists() else ""
    detailed_desc = (family_dir / "drawings" / "detailed_description_of_drawings.md").read_text(encoding="utf-8") if (family_dir / "drawings" / "detailed_description_of_drawings.md").exists() else ""

    # Use whichever document is more complete (longer) as the main spec
    main_text = prov_text if len(prov_text) > len(spec_text) else spec_text

    # Count drawing sheets
    drawings_dir = family_dir / "drawings" / "images"
    png_files = sorted([f for f in drawings_dir.iterdir() if f.suffix == ".png"]) if drawings_dir.exists() else []
    num_drawing_sheets = len(png_files)

    # Extract abstract
    abstract = extract_abstract(prov_text) or extract_abstract(spec_text)
    if not abstract:
        abstract = f"A system and method for {title.lower()[:200]}."

    # Truncate abstract to 150 words
    words = abstract.split()
    if len(words) > 150:
        abstract = " ".join(words[:150]) + "..."

    # === BUILD THE DOCUMENT ===
    doc = setup_doc()

    # --- Cover / Title ---
    add_para(doc, "NONPROVISIONAL UTILITY PATENT APPLICATION",
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    doc.add_paragraph()

    add_para(doc, f"Title of Invention: {title}", bold=True, size=12)
    doc.add_paragraph()
    add_para(doc, f"Applicant: The Wayfinder Trust")
    add_para(doc, f"Inventor: Jonathan Luethke")
    add_para(doc, f"Total Number of Drawing Sheets: {num_drawing_sheets}")
    add_para(doc, f"Total Number of Claims: See Claims section")

    doc.add_page_break()

    # Reset paragraph counter for specification numbering
    reset_para_counter()

    # --- Cross-Reference to Related Applications ---
    cross_ref_match = re.search(
        r"##\s*(?:CROSS-REFERENCE TO RELATED APPLICATIONS|Cross-Reference)\s*\n(.*?)(?=\n##)",
        main_text, re.DOTALL | re.IGNORECASE
    )
    if cross_ref_match:
        add_heading(doc, "CROSS-REFERENCE TO RELATED APPLICATIONS", level=1)
        add_para(doc, cross_ref_match.group(1).strip(), numbered=True)
        doc.add_paragraph()

    # --- Field of the Invention ---
    field_match = re.search(
        r"##\s*(?:\d+\.\s*)?FIELD OF THE INVENTION\s*\n(.*?)(?=\n##)",
        main_text, re.DOTALL | re.IGNORECASE
    )
    if field_match:
        add_heading(doc, "FIELD OF THE INVENTION", level=1)
        add_para(doc, field_match.group(1).strip(), numbered=True)
        doc.add_paragraph()

    # --- Background ---
    bg_match = re.search(
        r"##\s*(?:\d+\.\s*)?BACKGROUND[^#]*?\n(.*?)(?=\n##\s*(?:\d+\.\s*)?(?:SUMMARY|Summary))",
        main_text, re.DOTALL | re.IGNORECASE
    )
    if bg_match:
        add_heading(doc, "BACKGROUND OF THE INVENTION", level=1)
        bg_text = bg_match.group(1).strip()
        md_to_docx(doc, bg_text, skip_title=False, numbered=True)
        doc.add_paragraph()

    # --- Summary ---
    summary_match = re.search(
        r"##\s*(?:\d+\.\s*)?SUMMARY[^#]*?\n(.*?)(?=\n##\s*(?:\d+\.\s*)?(?:SYSTEM|DETAILED|Brief|Drawing))",
        main_text, re.DOTALL | re.IGNORECASE
    )
    if summary_match:
        add_heading(doc, "SUMMARY OF THE INVENTION", level=1)
        add_para(doc, summary_match.group(1).strip(), numbered=True)
        doc.add_paragraph()

    # --- Brief Description of the Drawings ---
    # Extract from detailed_description_of_drawings.md
    if detailed_desc:
        brief_match = re.search(
            r"##\s*Brief Description of the Drawings\s*\n(.*?)(?=\n##)",
            detailed_desc, re.DOTALL | re.IGNORECASE
        )
        if brief_match:
            add_heading(doc, "BRIEF DESCRIPTION OF THE DRAWINGS", level=1)
            for line in brief_match.group(1).strip().split("\n"):
                line = line.strip()
                if line:
                    add_para(doc, line, numbered=True)
            doc.add_paragraph()

    # --- Detailed Description ---
    add_heading(doc, "DETAILED DESCRIPTION OF THE INVENTION", level=1)

    # Insert the detailed description of drawings first
    if detailed_desc:
        dd_match = re.search(
            r"##\s*Detailed Description of the Drawings\s*\n(.*?)$",
            detailed_desc, re.DOTALL | re.IGNORECASE
        )
        if dd_match:
            for para in dd_match.group(1).strip().split("\n\n"):
                para = para.strip()
                if para:
                    add_para(doc, para, numbered=True)

    doc.add_paragraph()

    # Now add the main specification body
    # Skip sections we already added (title, cross-ref, field, background, summary, figures, abstract, claims)
    skip_patterns = [
        "cross-reference", "field of the invention", "background",
        "summary", "abstract", "claims", "figure", "drawing",
        "provisional patent", "invention specification",
    ]

    # Find the main body (everything after summary, before claims/abstract/figures)
    body_match = re.search(
        r"##\s*(?:\d+\.\s*)?(?:SYSTEM|DETAILED|4\.|Architecture)\s*(.*?)(?=\n##\s*(?:CLAIMS|ABSTRACT|FIGURE|BRIEF|DRAWING|CROSS-REFERENCE|\d+\.\s*DRAWING))",
        main_text, re.DOTALL | re.IGNORECASE
    )
    if body_match:
        body = body_match.group(0).strip()
        md_to_docx(doc, body, skip_title=False, skip_sections=[
            "claims", "abstract", "figure", "drawing"
        ], numbered=True)

    doc.add_page_break()

    # --- Claims ---
    add_heading(doc, "CLAIMS", level=1)
    add_para(doc, "What is claimed is:", italic=True)
    doc.add_paragraph()

    # Process claims - extract only actual claim text
    # Find all claims as numbered paragraphs
    claim_pattern = re.findall(
        r"\*\*(?:Claim\s+)?(\d+)\.?\*\*\.?\s*(.*?)(?=\*\*(?:Claim\s+)?\d+\.?\*\*|## CROSS-REFERENCE|## Cross-Reference|\*\*Total Claims|\Z)",
        claims_text, re.DOTALL
    )

    for num_str, body in claim_pattern:
        body = body.strip()
        # Remove section headers, horizontal rules, and markdown artifacts
        body = re.sub(r"^---\s*$", "", body, flags=re.MULTILINE)
        body = re.sub(r"^#{1,4}\s+.*$", "", body, flags=re.MULTILINE)
        body = re.sub(r"^\*\*[^*]+\*\*\s*$", "", body, flags=re.MULTILINE)
        body = re.sub(r"^\*[^*]+\*\s*$", "", body, flags=re.MULTILINE)
        body = body.strip()
        if not body:
            continue
        # Remove inline markdown bold
        body = re.sub(r"\*\*([^*]+)\*\*", r"\1", body)
        # Ensure ends with period
        if not body.endswith("."):
            body = body.rstrip() + "."
        # Collapse multiple whitespace/newlines into single space
        body = re.sub(r"\s+", " ", body).strip()

        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.27)
        run = p.add_run(f"{num_str}. {body}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    doc.add_page_break()

    # --- Abstract ---
    add_heading(doc, "ABSTRACT", level=1)
    add_para(doc, abstract)

    doc.add_page_break()

    # --- Drawing Sheets ---
    add_heading(doc, "DRAWINGS", level=1)
    add_para(doc, f"Total Number of Drawing Sheets: {num_drawing_sheets}", bold=True)
    doc.add_paragraph()

    for png in png_files:
        num_match = re.search(r"fig_(\d+)", png.stem)
        fig_num = int(num_match.group(1)) if num_match else 0
        fig_label = f"FIG. {fig_num}"

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(fig_label)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(png), width=Inches(5.5))

        if png != png_files[-1]:
            doc.add_page_break()

    # --- Save ---
    os.makedirs(output_dir, exist_ok=True)
    output_path = Path(output_dir) / f"{family_name}_NonProvisional.docx"
    doc.save(str(output_path))
    size_kb = os.path.getsize(output_path) // 1024
    print(f"  {output_path.name} ({size_kb} KB)")
    print(f"    Title: {title}")
    print(f"    Drawing Sheets: {num_drawing_sheets}")

    return output_path


def build_drawings_docx(family_name, output_dir):
    """Build a drawings-only DOCX (one figure per page)."""
    family_dir = PATENTS_DIR / family_name
    drawings_dir = family_dir / "drawings" / "images"
    if not drawings_dir.exists():
        return None

    doc = setup_doc()
    png_files = sorted([f for f in drawings_dir.iterdir() if f.suffix == ".png"])

    for idx, png in enumerate(png_files):
        num_match = re.search(r"fig_(\d+)", png.stem)
        fig_num = int(num_match.group(1)) if num_match else idx + 1

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"FIG. {fig_num}")
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(png), width=Inches(6.0))

        if idx < len(png_files) - 1:
            doc.add_page_break()

    os.makedirs(output_dir, exist_ok=True)
    output_path = Path(output_dir) / f"{family_name}_Drawings.docx"
    doc.save(str(output_path))
    print(f"  {output_path.name} ({os.path.getsize(output_path) // 1024} KB) - {len(png_files)} sheets")
    return output_path


def main():
    parser = argparse.ArgumentParser(
        description="Build USPTO non-provisional utility patent DOCX files",
    )
    parser.add_argument("--family", "-f", required=True, help="Family name or 'all'")
    parser.add_argument("--output", "-o", default=str(PATENTS_DIR / "filing_ready"),
                        help="Output directory")

    args = parser.parse_args()
    families = FAMILIES if args.family.lower() == "all" else [args.family]

    for family in families:
        print(f"\n=== {family} ===")
        build_nonprovisional_docx(family, args.output)
        build_drawings_docx(family, args.output)

    print(f"\nAll files in: {args.output}")


if __name__ == "__main__":
    main()
